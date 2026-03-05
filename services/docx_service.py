import os
import tempfile
from docx import Document
from docx.shared import Inches


def build_docx(cv: dict) -> str:
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    doc = Document()

    photo = cv.get("photo_path")
    if photo:
        try:
            doc.add_picture(photo, width=Inches(1.2))
        except Exception:
            pass

    doc.add_heading(cv.get("name", ""), level=0)
    if cv.get("title"):
        doc.add_paragraph(cv["title"])

    contact = []
    for k in ["email", "phone", "city"]:
        v = cv.get(k, "")
        if v:
            contact.append(v)
    if contact:
        doc.add_paragraph(" • ".join(contact))

    links = [cv.get("linkedin", ""), cv.get("github", ""), cv.get("portfolio", "")]
    links = [l for l in links if l]
    if links:
        doc.add_paragraph(" | ".join(links))

    lang = cv.get("lang", "pt")
    headings = {
        "Resumo": "Resumo" if lang=="pt" else "Summary",
        "Objetivo": "Objetivo" if lang=="pt" else "Objective",
        "Habilidades": "Habilidades" if lang=="pt" else "Skills",
        "Experiência": "Experiência" if lang=="pt" else "Experience",
        "Formação": "Formação" if lang=="pt" else "Education",
        "Certificações": "Certificações" if lang=="pt" else "Certifications",
        "Idiomas": "Idiomas" if lang=="pt" else "Languages",
    }

    def add_section(title: str):
        doc.add_heading(title, level=1)

    if cv.get("summary"):
        add_section(headings["Resumo"])
        doc.add_paragraph(cv["summary"])

    if cv.get("objective"):
        add_section(headings["Objetivo"])
        doc.add_paragraph(cv["objective"])

    skills = cv.get("skills", [])
    if skills:
        add_section(headings["Habilidades"])
        doc.add_paragraph(", ".join(skills))

    exp = cv.get("experience", [])
    if exp:
        add_section(headings["Experiência"])
        for line in exp:
            if line.startswith("•"):
                doc.add_paragraph(line.replace("•", "").strip(), style="List Bullet")
            else:
                p = doc.add_paragraph(line)
                if p.runs:
                    p.runs[0].bold = True

    edu = cv.get("education", [])
    if edu:
        add_section(headings["Formação"])
        for line in edu:
            doc.add_paragraph(line, style="List Bullet")

    certs = cv.get("certs", [])
    if certs:
        add_section(headings["Certificações"])
        for line in certs:
            doc.add_paragraph(line, style="List Bullet")

    langs = cv.get("languages", [])
    if langs:
        add_section(headings["Idiomas"])
        for line in langs:
            doc.add_paragraph(line, style="List Bullet")

    doc.save(path)
    return path